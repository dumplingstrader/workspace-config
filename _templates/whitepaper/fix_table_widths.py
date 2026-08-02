#!/usr/bin/env python3
"""
Fix table column widths and add borders in DOCX documents.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_table_borders(table):
    """Add professional borders to table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Create table borders element
    tblBorders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="12" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="12" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="12" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="12" w:space="0" w:color="auto"/>
    </w:tblBorders>
    '''
    
    # Remove existing borders if present
    existing_borders = tblPr.xpath("./w:tblBorders")
    for border in existing_borders:
        tblPr.remove(border)
    
    # Add new borders
    tblPr.append(parse_xml(tblBorders_xml))

def fix_table_widths(input_file, output_file):
    doc = Document(input_file)
    tables_fixed = 0
    table_width = Inches(6.0)
    
    for table_idx, table in enumerate(doc.tables):
        num_cols = len(table.rows[0].cells) if table.rows else 0
        if num_cols == 0:
            continue
        
        col_width = table_width / num_cols
        width_twips = int(col_width / 12700 * 20)
        
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                # Remove existing width elements
                for tcW in tcPr.xpath("./w:tcW"):
                    tcPr.remove(tcW)
                
                # Correctly formatted XML for the width element
                tcW_xml = f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>'
                tcPr.append(parse_xml(tcW_xml))
        
        # Add table borders
        add_table_borders(table)
        
        tables_fixed += 1
        print(f"  ✓ Table {table_idx + 1}: {num_cols} columns with borders")
    
    doc.save(output_file)
    return tables_fixed

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python fix_table_widths.py <input_docx> <output_docx>")
        sys.exit(1)
    
    try:
        num_fixed = fix_table_widths(sys.argv[1], sys.argv[2])
        print(f"\n✅ Complete! Fixed {num_fixed} tables")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
